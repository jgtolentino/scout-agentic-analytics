#!/usr/bin/env python3
"""
Iska - Enterprise Documentation & Asset Intelligence Agent v2.0
Core ingestion script with Supabase integration

This script handles:
- Web scraping of assets, SKUs, and documents
- PDF/DOCX/manual parsing and enrichment
- Change detection and incremental updates
- QA validation with Caca integration
- Audit logging and agent orchestration
- Semantic search embedding generation
"""

import asyncio
import json
import logging
import os
import hashlib
import time
from datetime import datetime, timezone
from pathlib import Path
from typing import Dict, List, Optional, Any, Tuple
from dataclasses import dataclass, asdict
from urllib.parse import urljoin, urlparse
import uuid

# Core dependencies
import aiohttp
import yaml
from supabase import create_client, Client
from openai import OpenAI
import PyPDF2
import docx
from bs4 import BeautifulSoup
import requests
from watchdog.observers import Observer
from watchdog.events import FileSystemEventHandler

# Configuration
CONFIG_FILE = "/Users/tbwa/agents/iska.yaml"
AUDIT_LOG_FILE = "/Users/tbwa/agents/logs/iska_audit.json"
CLAUDE_MD_PATH = "/Users/tbwa/CLAUDE.md"

# Logging setup
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s',
    handlers=[
        logging.FileHandler('/Users/tbwa/agents/logs/iska.log'),
        logging.StreamHandler()
    ]
)
logger = logging.getLogger(__name__)

@dataclass
class Document:
    """Document metadata and content structure"""
    id: str
    title: str
    content: str
    source: str
    source_type: str
    document_type: str
    file_path: Optional[str] = None
    url: Optional[str] = None
    checksum: Optional[str] = None
    file_size: Optional[int] = None
    created_at: Optional[datetime] = None
    updated_at: Optional[datetime] = None
    metadata: Optional[Dict[str, Any]] = None
    embedding: Optional[List[float]] = None
    qa_status: str = "pending"
    qa_errors: Optional[List[str]] = None

@dataclass
class Asset:
    """Asset metadata structure"""
    id: str
    asset_name: str
    asset_type: str
    asset_url: str
    brand: Optional[str] = None
    category: Optional[str] = None
    tags: Optional[List[str]] = None
    file_size: Optional[int] = None
    checksum: Optional[str] = None
    created_at: Optional[datetime] = None
    updated_at: Optional[datetime] = None
    metadata: Optional[Dict[str, Any]] = None

@dataclass
class AuditEntry:
    """Audit log entry structure"""
    timestamp: datetime
    source_type: str
    source_url: str
    document_type: str
    action: str
    agent_trigger: str
    qa_status: str
    error_message: Optional[str] = None
    processing_time: Optional[float] = None
    file_size: Optional[int] = None
    checksum: Optional[str] = None

class IskaIngestor:
    """Main Iska ingestion engine"""
    
    def __init__(self, config_path: str = CONFIG_FILE):
        self.config = self._load_config(config_path)
        self.supabase = self._init_supabase()
        self.openai = self._init_openai()
        self.session = None
        self.audit_log = []
        
    def _load_config(self, config_path: str) -> Dict[str, Any]:
        """Load Iska configuration from YAML file"""
        try:
            with open(config_path, 'r') as f:
                return yaml.safe_load(f)
        except Exception as e:
            logger.error(f"Failed to load config: {e}")
            raise

    def _init_supabase(self) -> Client:
        """Initialize Supabase client"""
        try:
            url = os.getenv('SUPABASE_URL')
            key = os.getenv('SUPABASE_SERVICE_ROLE_KEY')
            if not url or not key:
                raise ValueError("Supabase credentials not found in environment")
            return create_client(url, key)
        except Exception as e:
            logger.error(f"Failed to initialize Supabase: {e}")
            raise

    def _init_openai(self) -> OpenAI:
        """Initialize OpenAI client for embeddings"""
        try:
            api_key = os.getenv('OPENAI_API_KEY')
            if not api_key:
                raise ValueError("OpenAI API key not found in environment")
            return OpenAI(api_key=api_key)
        except Exception as e:
            logger.error(f"Failed to initialize OpenAI: {e}")
            raise

    async def _init_session(self) -> aiohttp.ClientSession:
        """Initialize aiohttp session"""
        if not self.session:
            timeout = aiohttp.ClientTimeout(total=30)
            self.session = aiohttp.ClientSession(timeout=timeout)
        return self.session

    def _calculate_checksum(self, content: str) -> str:
        """Calculate SHA256 checksum of content"""
        return hashlib.sha256(content.encode()).hexdigest()

    def _log_audit_entry(self, entry: AuditEntry):
        """Log audit entry to file and memory"""
        self.audit_log.append(entry)
        
        # Write to audit log file
        try:
            os.makedirs(os.path.dirname(AUDIT_LOG_FILE), exist_ok=True)
            with open(AUDIT_LOG_FILE, 'a') as f:
                f.write(json.dumps(asdict(entry), default=str) + '\n')
        except Exception as e:
            logger.error(f"Failed to write audit log: {e}")

    async def scrape_web_sources(self) -> List[Document]:
        """Scrape configured web sources for documents and assets"""
        documents = []
        session = await self._init_session()
        
        for source in self.config.get('ingestion_sources', {}).get('web_scraping', []):
            try:
                start_time = time.time()
                logger.info(f"Scraping {source['category']} from {source['url']}")
                
                async with session.get(source['url']) as response:
                    if response.status == 200:
                        html = await response.text()
                        docs = self._parse_web_content(html, source)
                        documents.extend(docs)
                        
                        # Log successful scrape
                        self._log_audit_entry(AuditEntry(
                            timestamp=datetime.now(timezone.utc),
                            source_type="web_scraping",
                            source_url=source['url'],
                            document_type=source['category'],
                            action="scrape_success",
                            agent_trigger="scheduled",
                            qa_status="pending",
                            processing_time=time.time() - start_time
                        ))
                    else:
                        logger.error(f"Failed to scrape {source['url']}: {response.status}")
                        
            except Exception as e:
                logger.error(f"Error scraping {source['url']}: {e}")
                self._log_audit_entry(AuditEntry(
                    timestamp=datetime.now(timezone.utc),
                    source_type="web_scraping",
                    source_url=source['url'],
                    document_type=source['category'],
                    action="scrape_error",
                    agent_trigger="scheduled",
                    qa_status="failed",
                    error_message=str(e),
                    processing_time=time.time() - start_time
                ))
        
        return documents

    def _parse_web_content(self, html: str, source: Dict[str, Any]) -> List[Document]:
        """Parse HTML content using configured selectors"""
        documents = []
        soup = BeautifulSoup(html, 'html.parser')
        
        # Extract based on configured selectors
        selectors = source.get('selectors', {})
        
        if 'asset_container' in selectors:
            # Asset extraction
            containers = soup.select(selectors['asset_container'])
            for container in containers:
                try:
                    doc = self._extract_asset_from_container(container, selectors, source)
                    if doc:
                        documents.append(doc)
                except Exception as e:
                    logger.error(f"Error extracting asset: {e}")
                    
        elif 'doc_container' in selectors:
            # Document extraction
            containers = soup.select(selectors['doc_container'])
            for container in containers:
                try:
                    doc = self._extract_document_from_container(container, selectors, source)
                    if doc:
                        documents.append(doc)
                except Exception as e:
                    logger.error(f"Error extracting document: {e}")
        
        return documents

    def _extract_asset_from_container(self, container, selectors: Dict[str, str], source: Dict[str, Any]) -> Optional[Document]:
        """Extract asset information from HTML container"""
        try:
            title_elem = container.select_one(selectors.get('asset_name', ''))
            type_elem = container.select_one(selectors.get('asset_type', ''))
            url_elem = container.select_one(selectors.get('asset_url', ''))
            
            if not all([title_elem, type_elem, url_elem]):
                return None
                
            title = title_elem.get_text(strip=True)
            asset_type = type_elem.get_text(strip=True)
            asset_url = url_elem.get('href') or url_elem.get('src')
            
            if not asset_url:
                return None
                
            # Make absolute URL
            if not asset_url.startswith('http'):
                asset_url = urljoin(source['url'], asset_url)
            
            return Document(
                id=str(uuid.uuid4()),
                title=title,
                content=f"Asset: {title} ({asset_type})",
                source=source['url'],
                source_type="web_scraping",
                document_type="asset",
                url=asset_url,
                created_at=datetime.now(timezone.utc),
                metadata={
                    'asset_type': asset_type,
                    'category': source['category'],
                    'original_url': source['url']
                }
            )
            
        except Exception as e:
            logger.error(f"Error extracting asset: {e}")
            return None

    def _extract_document_from_container(self, container, selectors: Dict[str, str], source: Dict[str, Any]) -> Optional[Document]:
        """Extract document information from HTML container"""
        try:
            title_elem = container.select_one(selectors.get('doc_title', ''))
            url_elem = container.select_one(selectors.get('doc_url', ''))
            type_elem = container.select_one(selectors.get('doc_type', ''))
            
            if not all([title_elem, url_elem]):
                return None
                
            title = title_elem.get_text(strip=True)
            doc_url = url_elem.get('href')
            doc_type = type_elem.get_text(strip=True) if type_elem else 'document'
            
            if not doc_url:
                return None
                
            # Make absolute URL
            if not doc_url.startswith('http'):
                doc_url = urljoin(source['url'], doc_url)
            
            return Document(
                id=str(uuid.uuid4()),
                title=title,
                content=f"Document: {title}",
                source=source['url'],
                source_type="web_scraping",
                document_type=doc_type,
                url=doc_url,
                created_at=datetime.now(timezone.utc),
                metadata={
                    'category': source['category'],
                    'original_url': source['url']
                }
            )
            
        except Exception as e:
            logger.error(f"Error extracting document: {e}")
            return None

    def ingest_local_documents(self) -> List[Document]:
        """Ingest documents from local filesystem"""
        documents = []
        
        for source in self.config.get('ingestion_sources', {}).get('document_sources', []):
            try:
                path = Path(source['path'])
                if not path.exists():
                    logger.warning(f"Path does not exist: {path}")
                    continue
                
                extensions = source.get('extensions', [])
                
                # Find all matching files
                for ext in extensions:
                    pattern = f"**/*{ext}"
                    for file_path in path.glob(pattern):
                        if file_path.is_file():
                            try:
                                doc = self._process_local_file(file_path, source)
                                if doc:
                                    documents.append(doc)
                            except Exception as e:
                                logger.error(f"Error processing file {file_path}: {e}")
                                
            except Exception as e:
                logger.error(f"Error ingesting from {source['path']}: {e}")
        
        return documents

    def _process_local_file(self, file_path: Path, source: Dict[str, Any]) -> Optional[Document]:
        """Process a single local file"""
        try:
            start_time = time.time()
            
            # Get file stats
            stat = file_path.stat()
            file_size = stat.st_size
            
            # Extract content based on file type
            content = ""
            if file_path.suffix.lower() == '.pdf':
                content = self._extract_pdf_content(file_path)
            elif file_path.suffix.lower() == '.docx':
                content = self._extract_docx_content(file_path)
            elif file_path.suffix.lower() in ['.md', '.txt']:
                content = self._extract_text_content(file_path)
            else:
                logger.warning(f"Unsupported file type: {file_path.suffix}")
                return None
            
            if not content.strip():
                logger.warning(f"No content extracted from {file_path}")
                return None
            
            # Calculate checksum
            checksum = self._calculate_checksum(content)
            
            # Create document
            doc = Document(
                id=str(uuid.uuid4()),
                title=file_path.stem,
                content=content,
                source=str(file_path),
                source_type="local_file",
                document_type=source['type'],
                file_path=str(file_path),
                checksum=checksum,
                file_size=file_size,
                created_at=datetime.fromtimestamp(stat.st_ctime, timezone.utc),
                updated_at=datetime.fromtimestamp(stat.st_mtime, timezone.utc),
                metadata={
                    'file_extension': file_path.suffix,
                    'source_type': source['type']
                }
            )
            
            # Log successful processing
            self._log_audit_entry(AuditEntry(
                timestamp=datetime.now(timezone.utc),
                source_type="local_file",
                source_url=str(file_path),
                document_type=source['type'],
                action="file_processed",
                agent_trigger="scheduled",
                qa_status="pending",
                processing_time=time.time() - start_time,
                file_size=file_size,
                checksum=checksum
            ))
            
            return doc
            
        except Exception as e:
            logger.error(f"Error processing file {file_path}: {e}")
            self._log_audit_entry(AuditEntry(
                timestamp=datetime.now(timezone.utc),
                source_type="local_file",
                source_url=str(file_path),
                document_type=source.get('type', 'unknown'),
                action="file_error",
                agent_trigger="scheduled",
                qa_status="failed",
                error_message=str(e),
                processing_time=time.time() - start_time
            ))
            return None

    def _extract_pdf_content(self, file_path: Path) -> str:
        """Extract text content from PDF file"""
        try:
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                content = ""
                for page in pdf_reader.pages:
                    content += page.extract_text() + "\n"
                return content.strip()
        except Exception as e:
            logger.error(f"Error extracting PDF content from {file_path}: {e}")
            return ""

    def _extract_docx_content(self, file_path: Path) -> str:
        """Extract text content from DOCX file"""
        try:
            doc = docx.Document(file_path)
            content = ""
            for paragraph in doc.paragraphs:
                content += paragraph.text + "\n"
            return content.strip()
        except Exception as e:
            logger.error(f"Error extracting DOCX content from {file_path}: {e}")
            return ""

    def _extract_text_content(self, file_path: Path) -> str:
        """Extract content from text-based files"""
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                return file.read()
        except Exception as e:
            logger.error(f"Error extracting text content from {file_path}: {e}")
            return ""

    async def qa_validation(self, documents: List[Document]) -> List[Document]:
        """Perform QA validation on documents"""
        validated_docs = []
        
        for doc in documents:
            try:
                # Perform validation checks
                errors = []
                
                # Check required fields
                if not doc.title or not doc.content:
                    errors.append("Missing required fields: title or content")
                
                # Check content length
                content_length = len(doc.content)
                min_length = self.config.get('qa_workflow', {}).get('validation_rules', [{}])[0].get('min_length', 100)
                max_length = self.config.get('qa_workflow', {}).get('validation_rules', [{}])[0].get('max_length', 100000)
                
                if content_length < min_length:
                    errors.append(f"Content too short: {content_length} < {min_length}")
                elif content_length > max_length:
                    errors.append(f"Content too long: {content_length} > {max_length}")
                
                # Check for duplicates (simplified check)
                if doc.checksum:
                    existing = await self._check_existing_document(doc.checksum)
                    if existing:
                        errors.append(f"Duplicate document found: {existing['id']}")
                
                # Update QA status
                if errors:
                    doc.qa_status = "failed"
                    doc.qa_errors = errors
                    logger.warning(f"QA validation failed for {doc.title}: {errors}")
                else:
                    doc.qa_status = "passed"
                    logger.info(f"QA validation passed for {doc.title}")
                
                validated_docs.append(doc)
                
            except Exception as e:
                logger.error(f"Error during QA validation for {doc.title}: {e}")
                doc.qa_status = "error"
                doc.qa_errors = [str(e)]
                validated_docs.append(doc)
        
        return validated_docs

    async def _check_existing_document(self, checksum: str) -> Optional[Dict[str, Any]]:
        """Check if document already exists in database"""
        try:
            result = self.supabase.table('agent_repository.documents').select('id').eq('checksum', checksum).execute()
            return result.data[0] if result.data else None
        except Exception as e:
            logger.error(f"Error checking existing document: {e}")
            return None

    async def generate_embeddings(self, documents: List[Document]) -> List[Document]:
        """Generate embeddings for semantic search"""
        for doc in documents:
            try:
                if doc.qa_status == "passed":
                    # Generate embedding
                    response = self.openai.embeddings.create(
                        model="text-embedding-3-small",
                        input=doc.content[:8000]  # Limit to model's input size
                    )
                    doc.embedding = response.data[0].embedding
                    logger.info(f"Generated embedding for {doc.title}")
                    
            except Exception as e:
                logger.error(f"Error generating embedding for {doc.title}: {e}")
        
        return documents

    async def store_documents(self, documents: List[Document]) -> List[Document]:
        """Store documents in Supabase database"""
        stored_docs = []
        
        for doc in documents:
            try:
                # Prepare document data
                doc_data = {
                    'id': doc.id,
                    'title': doc.title,
                    'content': doc.content,
                    'source': doc.source,
                    'source_type': doc.source_type,
                    'document_type': doc.document_type,
                    'file_path': doc.file_path,
                    'url': doc.url,
                    'checksum': doc.checksum,
                    'file_size': doc.file_size,
                    'created_at': doc.created_at.isoformat() if doc.created_at else None,
                    'updated_at': doc.updated_at.isoformat() if doc.updated_at else None,
                    'metadata': doc.metadata,
                    'qa_status': doc.qa_status,
                    'qa_errors': doc.qa_errors
                }
                
                # Store in documents table
                result = self.supabase.table('agent_repository.documents').upsert(doc_data).execute()
                
                # Store embedding if available
                if doc.embedding:
                    embedding_data = {
                        'id': str(uuid.uuid4()),
                        'document_id': doc.id,
                        'embedding': doc.embedding,
                        'created_at': datetime.now(timezone.utc).isoformat()
                    }
                    self.supabase.table('agent_repository.embeddings').upsert(embedding_data).execute()
                
                stored_docs.append(doc)
                logger.info(f"Stored document: {doc.title}")
                
            except Exception as e:
                logger.error(f"Error storing document {doc.title}: {e}")
        
        return stored_docs

    async def update_knowledge_base(self, documents: List[Document]):
        """Update CLAUDE.md and other knowledge base files"""
        try:
            # Update CLAUDE.md with new document references
            if self.config.get('knowledge_base', {}).get('claude_md_update'):
                await self._update_claude_md(documents)
            
            # Update SOP directory
            sop_docs = [doc for doc in documents if doc.document_type == 'SOPs']
            if sop_docs:
                await self._update_sop_directory(sop_docs)
            
            logger.info(f"Updated knowledge base with {len(documents)} documents")
            
        except Exception as e:
            logger.error(f"Error updating knowledge base: {e}")

    async def _update_claude_md(self, documents: List[Document]):
        """Update CLAUDE.md with new document references"""
        try:
            # Read current CLAUDE.md
            with open(CLAUDE_MD_PATH, 'r') as f:
                content = f.read()
            
            # Find or create Iska section
            iska_section = "\n## Iska Agent - Document Intelligence\n\n"
            iska_section += "### Recently Ingested Documents\n"
            
            for doc in documents:
                if doc.qa_status == "passed":
                    iska_section += f"- **{doc.title}** ({doc.document_type})\n"
                    iska_section += f"  - Source: {doc.source}\n"
                    iska_section += f"  - Updated: {doc.updated_at or doc.created_at}\n"
                    if doc.url:
                        iska_section += f"  - URL: {doc.url}\n"
                    iska_section += "\n"
            
            # Append or update section
            if "## Iska Agent - Document Intelligence" in content:
                # Replace existing section
                lines = content.split('\n')
                start_idx = None
                end_idx = None
                
                for i, line in enumerate(lines):
                    if "## Iska Agent - Document Intelligence" in line:
                        start_idx = i
                    elif start_idx is not None and line.startswith('## ') and i > start_idx:
                        end_idx = i
                        break
                
                if start_idx is not None:
                    if end_idx is not None:
                        lines[start_idx:end_idx] = iska_section.split('\n')
                    else:
                        lines[start_idx:] = iska_section.split('\n')
                    
                    content = '\n'.join(lines)
            else:
                # Append new section
                content += iska_section
            
            # Write back to file
            with open(CLAUDE_MD_PATH, 'w') as f:
                f.write(content)
            
            logger.info("Updated CLAUDE.md with new documents")
            
        except Exception as e:
            logger.error(f"Error updating CLAUDE.md: {e}")

    async def _update_sop_directory(self, sop_docs: List[Document]):
        """Update SOP directory with new documents"""
        try:
            sop_dir = Path(self.config.get('knowledge_base', {}).get('sop_directory', '/Users/tbwa/SOP/'))
            sop_dir.mkdir(exist_ok=True)
            
            for doc in sop_docs:
                if doc.qa_status == "passed":
                    # Create markdown file for SOP
                    filename = f"{doc.title.replace(' ', '_')}.md"
                    filepath = sop_dir / filename
                    
                    with open(filepath, 'w') as f:
                        f.write(f"# {doc.title}\n\n")
                        f.write(f"**Source**: {doc.source}\n")
                        f.write(f"**Type**: {doc.document_type}\n")
                        f.write(f"**Updated**: {doc.updated_at or doc.created_at}\n\n")
                        f.write("## Content\n\n")
                        f.write(doc.content)
            
            logger.info(f"Updated SOP directory with {len(sop_docs)} documents")
            
        except Exception as e:
            logger.error(f"Error updating SOP directory: {e}")

    async def notify_downstream_agents(self, documents: List[Document]):
        """Notify downstream agents about new documents"""
        try:
            routing_config = self.config.get('agent_routing', {})
            if not routing_config.get('enabled'):
                return
            
            for agent_config in routing_config.get('downstream_agents', []):
                agent_name = agent_config['name']
                trigger = agent_config['trigger']
                conditions = agent_config.get('conditions', [])
                
                # Check if any documents match the conditions
                matching_docs = []
                for doc in documents:
                    if doc.qa_status == "passed":
                        for condition in conditions:
                            if condition == "new_document":
                                matching_docs.append(doc)
                            elif condition == "document_update" and doc.updated_at:
                                matching_docs.append(doc)
                            elif condition == "claude_md_update" and doc.document_type == "system":
                                matching_docs.append(doc)
                
                if matching_docs:
                    await self._send_agent_notification(agent_name, trigger, matching_docs)
            
        except Exception as e:
            logger.error(f"Error notifying downstream agents: {e}")

    async def _send_agent_notification(self, agent_name: str, trigger: str, documents: List[Document]):
        """Send notification to a specific agent"""
        try:
            # Create notification payload
            notification = {
                'agent': agent_name,
                'trigger': trigger,
                'timestamp': datetime.now(timezone.utc).isoformat(),
                'documents': [
                    {
                        'id': doc.id,
                        'title': doc.title,
                        'document_type': doc.document_type,
                        'source': doc.source
                    } for doc in documents
                ],
                'count': len(documents)
            }
            
            # Store notification in database
            self.supabase.table('agent_repository.agent_notifications').insert(notification).execute()
            
            logger.info(f"Sent notification to {agent_name} for {len(documents)} documents")
            
        except Exception as e:
            logger.error(f"Error sending notification to {agent_name}: {e}")

    async def run_ingestion_cycle(self):
        """Run complete ingestion cycle"""
        try:
            logger.info("Starting Iska ingestion cycle")
            start_time = time.time()
            
            # Collect documents from all sources
            all_documents = []
            
            # Web scraping
            web_docs = await self.scrape_web_sources()
            all_documents.extend(web_docs)
            logger.info(f"Scraped {len(web_docs)} documents from web sources")
            
            # Local file ingestion
            local_docs = self.ingest_local_documents()
            all_documents.extend(local_docs)
            logger.info(f"Ingested {len(local_docs)} documents from local sources")
            
            # QA validation
            validated_docs = await self.qa_validation(all_documents)
            passed_docs = [doc for doc in validated_docs if doc.qa_status == "passed"]
            logger.info(f"QA validation: {len(passed_docs)} passed, {len(validated_docs) - len(passed_docs)} failed")
            
            # Generate embeddings
            embedded_docs = await self.generate_embeddings(passed_docs)
            logger.info(f"Generated embeddings for {len(embedded_docs)} documents")
            
            # Store in database
            stored_docs = await self.store_documents(embedded_docs)
            logger.info(f"Stored {len(stored_docs)} documents in database")
            
            # Update knowledge base
            await self.update_knowledge_base(stored_docs)
            
            # Notify downstream agents
            await self.notify_downstream_agents(stored_docs)
            
            # Final logging
            total_time = time.time() - start_time
            logger.info(f"Ingestion cycle completed in {total_time:.2f} seconds")
            
            # Store cycle summary
            summary = {
                'timestamp': datetime.now(timezone.utc).isoformat(),
                'total_documents': len(all_documents),
                'passed_qa': len(passed_docs),
                'stored_documents': len(stored_docs),
                'processing_time': total_time,
                'sources': {
                    'web_scraping': len(web_docs),
                    'local_files': len(local_docs)
                }
            }
            
            self.supabase.table('agent_repository.ingestion_cycles').insert(summary).execute()
            
        except Exception as e:
            logger.error(f"Error in ingestion cycle: {e}")
            raise
        finally:
            if self.session:
                await self.session.close()

async def main():
    """Main entry point for Iska ingestion"""
    try:
        # Initialize Iska ingestor
        iska = IskaIngestor()
        
        # Run ingestion cycle
        await iska.run_ingestion_cycle()
        
        logger.info("Iska ingestion completed successfully")
        
    except Exception as e:
        logger.error(f"Iska ingestion failed: {e}")
        raise

if __name__ == "__main__":
    asyncio.run(main())